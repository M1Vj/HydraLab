"""DOCX converter tests (branch 01-12).

Covers HL-EXPORT-06 (import metadata), HL-EXPORT-07 / HL-WRITE-20 (export to
outputs/ only + effective format applied), HL-EXPORT-08 / HL-EXPORT-09
(availability detection + persistence), HL-EXPORT-10 (failure no-mutation),
HL-WRITE-22 (traversal rejected + macro flagged), HL-WRITE-19 (LaTeX toolchain
detection) and HL-LIC-04 (permissive bundled converter).
"""
import hashlib
import shutil
import zipfile
from importlib.metadata import metadata
from pathlib import Path

import pytest
import pytest_asyncio
from sqlmodel import SQLModel
from sqlalchemy.ext.asyncio import create_async_engine
from sqlmodel.ext.asyncio.session import AsyncSession
from sqlalchemy.orm import sessionmaker

import docx

from hydra.database.repository import Repository
from hydra.services.docx import (
    ConverterAvailability,
    DocxPackageError,
    DocxService,
    PythonDocxAdapter,
    detect_latex_toolchain,
    scrub_secrets,
)
from hydra.services.docx.adapters import ConverterAdapter
from hydra.services.writing import global_defaults_from_settings, merge_format
from hydra.settings.toml_config import default_settings


# --- fixtures ---------------------------------------------------------------


@pytest_asyncio.fixture
async def session():
    engine = create_async_engine("sqlite+aiosqlite:///:memory:", future=True)
    async with engine.begin() as conn:
        await conn.run_sync(SQLModel.metadata.create_all)
    maker = sessionmaker(engine, class_=AsyncSession, expire_on_commit=False)
    async with maker() as db:
        yield db
    await engine.dispose()


def _benign_docx(path: Path, title: str = "Attention Is All You Need", author: str = "Vaswani") -> Path:
    document = docx.Document()
    document.add_heading("Introduction", level=1)
    document.add_paragraph("Transformers rely on attention.")
    document.core_properties.title = title
    document.core_properties.author = author
    document.save(str(path))
    return path


def _macro_docx(path: Path) -> Path:
    benign = path.with_suffix(".seed.docx")
    _benign_docx(benign)
    with zipfile.ZipFile(benign, "a") as archive:
        archive.writestr("word/vbaProject.bin", b"\x00MACRO-PAYLOAD")
    benign.rename(path)
    return path


def _traversal_docx(path: Path) -> Path:
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr(
            "[Content_Types].xml",
            '<?xml version="1.0"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"/>',
        )
        archive.writestr("../../escape.txt", "pwned")
    return path


class _FailingExportAdapter:
    name = "failing"

    def detect(self) -> ConverterAvailability:
        return ConverterAvailability(adapter=self.name, version="1.0", available=True)

    def import_docx(self, path, extract_root):  # pragma: no cover - not used
        raise RuntimeError("not used")

    def export(self, text, out_path, fmt):
        raise RuntimeError("simulated converter crash")


def _fmt(**overrides):
    defaults = global_defaults_from_settings(default_settings())
    return merge_format(defaults, overrides)


# --- HL-EXPORT-08 / 09 ------------------------------------------------------


def test_hl_export_08_no_converter_shows_setup_state_and_no_success(tmp_path):
    service = DocxService(adapters=[])  # simulate no local converter
    availability = service.detect()
    assert availability.available is False
    assert availability.setup_error

    (tmp_path / "writing" / "manuscripts" / "m").mkdir(parents=True)
    (tmp_path / "writing" / "manuscripts" / "m" / "main.md").write_text("# Title\n")
    result = service.export_manuscript(tmp_path, "m", "main.md", _fmt())
    assert result.status == "unavailable"
    assert result.output_path is None
    assert not (tmp_path / "outputs" / "manuscripts" / "m").exists()


@pytest.mark.asyncio
async def test_hl_export_09_availability_recorded_and_survives_restart(session):
    repo = Repository(session)
    availability = DocxService(adapters=[]).detect()
    await repo.record_docx_artifact(
        kind="availability",
        converter_adapter=availability.adapter,
        converter_version=availability.version,
        availability_status=availability.status,
        setup_error=availability.setup_error,
        status=availability.status,
    )
    latest = await repo.latest_docx_availability()
    assert latest is not None
    assert latest["availability_status"] == "unavailable"
    assert latest["setup_error"]


# --- HL-EXPORT-06 -----------------------------------------------------------


def test_hl_export_06_import_shows_metadata_without_changing_original(tmp_path):
    original = _benign_docx(tmp_path / "submission.docx")
    before = hashlib.sha256(original.read_bytes()).hexdigest()

    result = DocxService().import_docx(original)

    assert result.status == "success"
    assert result.metadata.get("title") == "Attention Is All You Need"
    assert result.metadata.get("author") == "Vaswani"
    assert "attention" in result.content.lower()
    after = hashlib.sha256(original.read_bytes()).hexdigest()
    assert before == after  # original unchanged on disk


# --- HL-EXPORT-07 / HL-WRITE-20 ---------------------------------------------


def test_hl_export_07_export_applies_effective_format_to_outputs_only(tmp_path):
    manuscript_dir = tmp_path / "writing" / "manuscripts" / "transformer-survey"
    manuscript_dir.mkdir(parents=True)
    source = manuscript_dir / "main.md"
    source.write_text("# Introduction\n\nTransformers use attention.\n", encoding="utf-8")
    source_hash = hashlib.sha256(source.read_bytes()).hexdigest()

    fmt = _fmt(heading_numbering=True, line_spacing=1.5)
    result = DocxService().export_manuscript(tmp_path, "transformer-survey", "main.md", fmt)

    assert result.status == "success"
    out_path = Path(result.output_path)
    assert out_path.exists()
    assert out_path.parent == tmp_path / "outputs" / "manuscripts" / "transformer-survey"
    # writing/manuscripts working source unchanged (HL-WRITE-20).
    assert hashlib.sha256(source.read_bytes()).hexdigest() == source_hash

    produced = docx.Document(str(out_path))
    headings = [p.text for p in produced.paragraphs if p.style.name.startswith("Heading")]
    assert any(h.startswith("1 ") for h in headings)  # numbered heading
    body = [p for p in produced.paragraphs if not p.style.name.startswith("Heading") and p.text]
    assert body and abs(float(body[0].paragraph_format.line_spacing) - 1.5) < 0.01


def test_export_scrubs_secrets_from_output(tmp_path):
    manuscript_dir = tmp_path / "writing" / "manuscripts" / "leaky"
    manuscript_dir.mkdir(parents=True)
    (manuscript_dir / "main.md").write_text("# Draft\n\nkey sk-abcdefgh12345678 here\n", encoding="utf-8")

    result = DocxService().export_manuscript(tmp_path, "leaky", "main.md", _fmt())
    text = "\n".join(p.text for p in docx.Document(result.output_path).paragraphs)
    assert "sk-abcdefgh12345678" not in text
    assert "redacted-secret" in text


def test_scrub_secrets_redacts_known_shapes():
    assert "redacted" in scrub_secrets("token ghp_" + "a" * 30)
    assert "AKIA" not in scrub_secrets("AKIA" + "A" * 16)


def test_scrub_secrets_preserves_ai_research_prose():
    # The bare ``ai-`` shape was removed: an AI-research manuscript must keep its
    # ordinary hyphenated prose intact instead of being corrupted to [redacted].
    prose = "This AI-generated summary of the ai-assisted pipeline is human-reviewed."
    assert scrub_secrets(prose) == prose


def test_export_preserves_ai_prose_in_output(tmp_path):
    manuscript_dir = tmp_path / "writing" / "manuscripts" / "ai-paper"
    manuscript_dir.mkdir(parents=True)
    (manuscript_dir / "main.md").write_text(
        "# Draft\n\nAn ai-generated draft with ai-assisted edits.\n", encoding="utf-8"
    )

    result = DocxService().export_manuscript(tmp_path, "ai-paper", "main.md", _fmt())
    text = "\n".join(p.text for p in docx.Document(result.output_path).paragraphs)
    assert "ai-generated" in text
    assert "ai-assisted" in text
    assert "redacted" not in text


# --- HL-WRITE-20 path containment (attacker-controlled names) ----------------


def test_export_rejects_source_relpath_traversal(tmp_path):
    manuscript_dir = tmp_path / "writing" / "manuscripts" / "m"
    manuscript_dir.mkdir(parents=True)
    (manuscript_dir / "main.md").write_text("# Title\n", encoding="utf-8")
    secret = tmp_path.parent / "traversal-secret.txt"
    secret.write_text("TOP-SECRET", encoding="utf-8")

    result = DocxService().export_manuscript(
        tmp_path, "m", "../../../traversal-secret.txt", _fmt()
    )

    assert result.status == "failed"
    assert "escape" in result.error_detail.lower()
    assert not (tmp_path / "outputs").exists()


def test_export_rejects_manuscript_name_traversal(tmp_path):
    (tmp_path / "writing" / "manuscripts").mkdir(parents=True)
    result = DocxService().export_manuscript(
        tmp_path, "../../../etc", "main.md", _fmt()
    )
    assert result.status == "failed"
    assert "escape" in result.error_detail.lower()


def test_export_output_name_cannot_escape_outputs_dir(tmp_path):
    manuscript_dir = tmp_path / "writing" / "manuscripts" / "m"
    manuscript_dir.mkdir(parents=True)
    (manuscript_dir / "main.md").write_text("# Title\n\nBody\n", encoding="utf-8")
    escape_target = tmp_path.parent / "evil.docx"
    escape_target.unlink(missing_ok=True)

    result = DocxService().export_manuscript(
        tmp_path, "m", "main.md", _fmt(), output_name="../../../../evil.docx"
    )

    # The directory components are stripped: output lands safely inside outputs/.
    assert result.status == "success"
    out_path = Path(result.output_path)
    assert out_path.parent == tmp_path / "outputs" / "manuscripts" / "m"
    assert not escape_target.exists()


# --- HL-EXPORT-10 -----------------------------------------------------------


def test_hl_export_10_converter_failure_leaves_prior_export_unchanged(tmp_path):
    manuscript_dir = tmp_path / "writing" / "manuscripts" / "m"
    manuscript_dir.mkdir(parents=True)
    (manuscript_dir / "main.md").write_text("# Title\n\nBody\n", encoding="utf-8")

    prior = tmp_path / "outputs" / "manuscripts" / "m" / "main.docx"
    prior.parent.mkdir(parents=True)
    prior.write_bytes(b"PRIOR-EXPORT-CONTENT")
    prior_hash = hashlib.sha256(prior.read_bytes()).hexdigest()

    service = DocxService(adapters=[_FailingExportAdapter()])
    result = service.export_manuscript(tmp_path, "m", "main.md", _fmt(), output_name="main.docx")

    assert result.status == "failed"
    assert "simulated converter crash" in result.error_detail
    assert prior.read_bytes() == b"PRIOR-EXPORT-CONTENT"
    assert hashlib.sha256(prior.read_bytes()).hexdigest() == prior_hash


# --- HL-WRITE-22 ------------------------------------------------------------


def test_hl_write_22_path_traversal_docx_is_rejected(tmp_path):
    evil = _traversal_docx(tmp_path / "evil.docx")
    escape_target = tmp_path.parent / "escape.txt"
    escape_target.unlink(missing_ok=True)

    result = DocxService().import_docx(evil)

    assert result.status == "rejected"
    assert "escape" in result.error_detail.lower() or "traversal" in result.error_detail.lower()
    # Nothing written outside the temp extraction dir.
    assert not escape_target.exists()
    assert not (tmp_path / "escape.txt").exists()


def test_hl_write_22_macro_docx_is_flagged_and_not_executed(tmp_path):
    macro = _macro_docx(tmp_path / "macro.docx")

    result = DocxService().import_docx(macro)

    assert result.status == "success"
    assert any("vbaProject.bin" in flag for flag in result.flagged_active_content)


def test_direct_extract_cleans_up_and_never_escapes(tmp_path):
    from hydra.services.docx.security import extract_docx_safely

    macro = _macro_docx(tmp_path / "macro.docx")
    dest = tmp_path / "extract"
    extraction = extract_docx_safely(macro, dest)
    assert extraction.flagged_active_content
    assert not (dest / "word" / "vbaProject.bin").exists()  # macro skipped, not extracted


# --- HL-WRITE-19 ------------------------------------------------------------


def test_hl_write_19_latex_toolchain_detection_reports_setup_state(monkeypatch):
    monkeypatch.setattr("hydra.services.docx.service.shutil.which", lambda binary: None)
    state = detect_latex_toolchain()
    assert state["available"] is False
    assert "TeX toolchain" in state["setup_error"]


# --- HL-LIC-04 --------------------------------------------------------------


def test_hl_lic_04_bundled_converter_is_permissive():
    # The bundled default adapter is python-docx, SPDX MIT.
    assert PythonDocxAdapter().name == "python-docx"
    license_text = (metadata("python-docx").get("License") or "").upper()
    classifiers = " ".join(metadata("python-docx").get_all("Classifier") or []).upper()
    assert "MIT" in license_text or "MIT LICENSE" in classifiers
    assert "AGPL" not in license_text and "AGPL" not in classifiers


def test_default_service_detects_bundled_converter(tmp_path):
    # With python-docx installed, the default service is available.
    availability = DocxService().detect()
    assert availability.available is True
    assert availability.adapter == "python-docx"
